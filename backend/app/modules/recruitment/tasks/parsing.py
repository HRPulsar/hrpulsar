"""Resume parsing task and text-extraction helpers (PDF / DOCX via S3).

Split from the former recruitment/tasks.py monolith (project-review #20).
Task names are pinned to the pre-split ``app.modules.recruitment.tasks.*``
namespace -- they are a public contract (beat schedule, queued messages,
the task_failure status map).
"""

import logging

from app.core.celery_app import celery

logger = logging.getLogger(__name__)


@celery.task(
    bind=True,
    max_retries=2,
    default_retry_delay=30,
    name="app.modules.recruitment.tasks.parse_resume_task",
)
def parse_resume_task(
    self,
    file_id: str | None = None,
    tenant_id: str | None = None,
    *,
    resume_id: str | None = None,
) -> dict:
    """Parse a CandidateFile resume: text extraction → LLM → persist.

    HRP-181 REDO Stage 3:

    - The argument is the canonical ``candidate_files.id``; ``resume_id``
      is accepted as a back-compat keyword alias and removed in Stage 5.
    - When ``candidate_id IS NULL`` (detached upload from the bulk modal)
      the task only updates ``parsed_data`` / ``raw_text`` /
      ``parse_status``. Candidate / Person rows stay untouched —
      ``finalize_candidates_from_parsed`` materialises the canonical
      ``Candidate`` later when the user clicks "Import".
    - When ``candidate_id`` is set (legacy single-file upload) the task
      additionally backfills ``Candidate.parsed_resume_jsonb`` if it was
      empty. **Person is never touched** — the Person backfill from the
      pre-Stage 3 codepath wrote fields the canonical model no longer
      reads.
    """
    import asyncio
    import uuid

    from sqlalchemy import create_engine
    from sqlalchemy.orm import Session

    from app.config import settings
    from app.modules.recruitment.models import Candidate, CandidateFile

    target_id = file_id or resume_id
    if not target_id:
        logger.error("parse_resume_task called without file_id/resume_id")
        return {"status": "error", "error": "file_id missing"}

    # E2E mode: skip the LLM parse entirely so the Playwright
    # ``_test/seed-parsed-files`` endpoint owns the ``parse_status`` /
    # ``parsed_data`` transitions without a Celery race overwriting them
    # with ``failed`` (S3 + LLM are unavailable in the GitHub Actions
    # runner).
    if settings.e2e_mode:
        logger.info("parse_resume_task skipped under E2E_MODE for file %s", target_id)
        return {"status": "skipped", "reason": "e2e_mode"}

    sync_url = settings.database_url.replace("+asyncpg", "+psycopg2")
    engine = create_engine(sync_url)

    try:
        with Session(engine) as db:
            resume = db.get(CandidateFile, uuid.UUID(target_id))
            if not resume:
                logger.error("CandidateFile %s not found", target_id)
                return {"status": "error", "error": "Resume not found"}

            resume.parse_status = "processing"
            db.commit()

            text = resume.raw_text or ""
            if not text and resume.file_id:
                text = _extract_text_from_s3(resume, settings)

            if not text:
                resume.parse_status = "failed"
                resume.parsed_data = {"error": "Could not extract text from file"}
                db.commit()
                return {"status": "failed", "error": "No text extracted"}

            from app.modules.recruitment.ai_service import parse_resume_text

            parsed = asyncio.run(parse_resume_text(text))

            resume.parsed_data = parsed
            resume.raw_text = text
            resume.parse_status = "completed"

            if resume.candidate_id is not None:
                candidate = db.get(Candidate, resume.candidate_id)
                if (
                    candidate is not None
                    and parsed
                    and candidate.parsed_resume_jsonb is None
                ):
                    candidate.parsed_resume_jsonb = parsed

            db.commit()
            logger.info("CandidateFile %s parsed successfully", target_id)
            return {"status": "completed", "file_id": target_id}

    except Exception as exc:
        logger.exception("parse_resume_task failed for %s", target_id)
        try:
            with Session(engine) as db:
                resume = db.get(CandidateFile, uuid.UUID(target_id))
                if resume:
                    resume.parse_status = "failed"
                    resume.parsed_data = {"error": str(exc)}
                    db.commit()
        except Exception:
            logger.exception("Failed to update resume status to failed")
        raise self.retry(exc=exc)
    finally:
        engine.dispose()


def _extract_text_from_s3(resume, settings) -> str:
    """Download file from S3 and extract text based on mime type."""
    from app.core.s3 import get_s3_client

    client = get_s3_client()
    if not client:
        return ""

    # Build S3 key from file record
    try:
        from sqlalchemy import create_engine
        from sqlalchemy.orm import Session

        from app.modules.storage.models import File

        sync_url = settings.database_url.replace("+asyncpg", "+psycopg2")
        engine = create_engine(sync_url)
        try:
            with Session(engine) as db:
                file_record = db.get(File, resume.file_id)
                if not file_record:
                    return ""
                path = file_record.path
        finally:
            engine.dispose()

        response = client.get_object(Bucket=settings.s3_bucket, Key=path)
        data = response["Body"].read()
    except Exception:  # noqa: BLE001 - unreadable source file -> empty text
        return ""

    mime = resume.mime_type.lower()

    if "pdf" in mime:
        return _extract_pdf_text(data)
    elif "docx" in mime or "document" in mime:
        return _extract_docx_text(data)
    elif "text" in mime or "rtf" in mime:
        return data.decode("utf-8", errors="ignore")

    return ""


def _extract_pdf_text(data: bytes) -> str:
    """Extract text from PDF bytes using PyPDF2."""
    import io

    from PyPDF2 import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text)
    return "\n\n".join(pages)


def _extract_docx_text(data: bytes) -> str:
    """Extract text from DOCX bytes using python-docx.

    HRP-345: many resume templates keep all content in tables (two-column
    layouts) or headers, where ``doc.paragraphs`` sees nothing — walk
    tables (recursively for nested ones) and section headers/footers too.
    """
    import io

    from docx import Document

    doc = Document(io.BytesIO(data))
    blocks: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]

    def _walk_table(table) -> None:
        # Merged cells surface once per spanned grid slot in ``row.cells``
        # — dedupe on the underlying XML element.
        seen_tc: set = set()
        for row in table.rows:
            cells = []
            for cell in row.cells:
                if cell._tc in seen_tc:
                    continue
                seen_tc.add(cell._tc)
                if cell.text.strip():
                    cells.append(cell.text.strip())
                for nested in cell.tables:
                    _walk_table(nested)
            if cells:
                blocks.append(" | ".join(cells))

    for table in doc.tables:
        _walk_table(table)

    for section in doc.sections:
        for para in (*section.header.paragraphs, *section.footer.paragraphs):
            if para.text.strip():
                blocks.append(para.text)

    return "\n\n".join(blocks)
