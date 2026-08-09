"""HRP-186: public token endpoint helpers.

Lives outside the user-authenticated service path so the tenant_id is
*derived* from the invite — never read from a request header. The rate
limiter is best-effort: it persists IP blocks into the DB so multiple
backend workers share the deny-list without needing a separate Redis
slot for it.
"""

from __future__ import annotations

import asyncio
import logging
import uuid
from datetime import datetime, timedelta, timezone
from typing import Any

from fastapi import status
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.errors import AppError
from app.modules.recruitment import audit_service
from app.modules.recruitment.manager_assessment_models import (
    PublicAssessmentIPBlock,
)
from app.modules.recruitment.manager_assessment_service import (
    INVITE_ACTIVE_STATUSES,
    assert_round_open,
    get_assessment,
    get_or_create_assessment,
    hash_token,
    list_vacancy_profile_competences,
    recompute_manager_score,
)
from app.modules.recruitment.models import (
    AssessmentInvite,
    Candidate,
    CandidateFile,
    CandidateVacancy,
    Vacancy,
)

log = logging.getLogger(__name__)


PUBLIC_RATE_LIMIT_INVALID_THRESHOLD = 5
PUBLIC_RATE_LIMIT_WINDOW_MINUTES = 5
PUBLIC_RATE_LIMIT_BLOCK_HOURS = 1

# HRP-359: share of *critical* competences that must be scored before a
# frictionless submit; below it the UI shows a "Submit anyway?" warning.
# The check is advisory (client-side modal), never a hard server block.
CRITICAL_SUBMIT_THRESHOLD = 0.5

# HRP-371: a browser cannot render a .docx, so pointing an <iframe> at the
# presigned URL left the Resume pane blank *and* pushed the file into the
# evaluator's Downloads folder on every page load. The server extracts the
# document's text instead and the page renders it as an ordinary preview;
# the raw file only leaves S3 when the evaluator asks for it.
RESUME_PREVIEW_MAX_BLOCKS = 400
RESUME_PREVIEW_MAX_BYTES = 15 * 1024 * 1024
#: Mime fragments we can turn into a text preview server-side.
_DOCX_MIME_HINTS = ("wordprocessingml", "docx")
_PLAINTEXT_MIME_HINTS = ("text/", "rtf")


# ---------------------------------------------------------------------------
# IP block helpers
# ---------------------------------------------------------------------------


async def is_ip_blocked(db: AsyncSession, ip: str) -> bool:
    if not ip:
        return False
    result = await db.execute(
        select(PublicAssessmentIPBlock)
        .where(PublicAssessmentIPBlock.ip == ip)
        .order_by(PublicAssessmentIPBlock.blocked_until.desc())
        .limit(1)
    )
    block = result.scalar_one_or_none()
    if block is None:
        return False
    return block.blocked_until > datetime.now(timezone.utc)


async def register_invalid_attempt(db: AsyncSession, ip: str) -> None:
    """Track a failed lookup and trip the deny-list if threshold is hit."""
    if not ip:
        return
    cutoff = datetime.now(timezone.utc) - timedelta(
        minutes=PUBLIC_RATE_LIMIT_WINDOW_MINUTES
    )
    result = await db.execute(
        select(PublicAssessmentIPBlock)
        .where(
            PublicAssessmentIPBlock.ip == ip,
            PublicAssessmentIPBlock.created_at >= cutoff,
        )
        .order_by(PublicAssessmentIPBlock.created_at.desc())
        .limit(1)
    )
    block = result.scalar_one_or_none()
    if block is not None:
        block.failures += 1
        if block.failures >= PUBLIC_RATE_LIMIT_INVALID_THRESHOLD:
            block.blocked_until = datetime.now(timezone.utc) + timedelta(
                hours=PUBLIC_RATE_LIMIT_BLOCK_HOURS
            )
            block.reason = "invalid_token_threshold"
        await db.commit()
        return
    block = PublicAssessmentIPBlock(
        ip=ip,
        reason="invalid_token",
        failures=1,
        blocked_until=datetime.now(timezone.utc),
    )
    db.add(block)
    await db.commit()


# ---------------------------------------------------------------------------
# Token resolution
# ---------------------------------------------------------------------------


async def resolve_invite_by_token(
    db: AsyncSession,
    token: str,
    *,
    ip: str | None = None,
) -> AssessmentInvite:
    """Return the invite that owns the token, or raise.

    Status codes mirror the spec:
    * ``410`` for revoked / expired / unknown — the API never confirms
      whether a token was once valid.
    * ``403`` is reserved for cross-vacancy / cross-candidate misuse
      with a known good token.
    """

    if ip and await is_ip_blocked(db, ip):
        raise AppError("too_many_failed_attempts", status.HTTP_429_TOO_MANY_REQUESTS)
    if not token:
        raise AppError("invalid_invite_link", status.HTTP_410_GONE)

    token_h = hash_token(token)
    result = await db.execute(
        select(AssessmentInvite).where(
            (AssessmentInvite.token_hash == token_h) | (AssessmentInvite.token == token)
        )
    )
    invite = result.scalar_one_or_none()
    if invite is None:
        await register_invalid_attempt(db, ip or "")
        raise AppError("invalid_invite_link", status.HTTP_410_GONE)

    if invite.revoked_at is not None:
        raise AppError("invitation_revoked", status.HTTP_410_GONE)
    if invite.purged_at is not None:
        raise AppError("invitation_expired", status.HTTP_410_GONE)
    if invite.status == "declined":
        # HRP-359 REDO: declining is terminal — the token dies with it.
        # Refreshing the stub page must not resurface the consent screen.
        raise AppError("invitation_declined", status.HTTP_410_GONE)
    if invite.expires_at < datetime.now(timezone.utc):
        # Only live invites flip to ``expired`` — a terminal status
        # (submitted / declined / revoked) must survive a late link visit,
        # otherwise a completed evaluation would show as expired and
        # extend_invite's submitted-guard could be bypassed (review [1]).
        if invite.status in INVITE_ACTIVE_STATUSES:
            invite.status = "expired"
            await db.commit()
        raise AppError("invitation_expired", status.HTTP_410_GONE)
    return invite


# ---------------------------------------------------------------------------
# Public actions
# ---------------------------------------------------------------------------


def _ensure_editable(invite: AssessmentInvite) -> None:
    """Reject edits on a submitted sheet when re-editing is disabled."""
    if invite.status == "submitted" and not invite.allow_reediting:
        raise AppError(
            "evaluation_submitted_reediting_disabled", status.HTTP_409_CONFLICT
        )


async def _mark_opened(db: AsyncSession, invite: AssessmentInvite) -> None:
    """HRP-358: following the email link flips ``pending`` → ``opened``."""
    if invite.status != "pending" and invite.opened_at is not None:
        return
    if invite.status == "pending":
        invite.status = "opened"
    if invite.opened_at is None:
        invite.opened_at = datetime.now(timezone.utc)
    await db.commit()
    await audit_service.record_event(
        db,
        tenant_id=invite.tenant_id,
        user_id=None,
        action="assessment_invite.opened",
        entity_type="assessment_invite",
        entity_id=invite.id,
    )


async def _assert_invite_round_open(db: AsyncSession, invite: AssessmentInvite) -> None:
    """HRP-376: the external sheet dies with its round.

    Completing or archiving a round makes every sheet on it read-only —
    the score endpoints already refuse through ``set_*_score``, but submit
    / notes / name went straight to the invite and would flip it (and the
    sheet) to ``submitted`` with no scores behind it. An invite whose
    round was archived rather than completed still holds a live token, so
    this is the only thing standing between that page and a write.
    """
    if invite.round_id is None:
        return
    await assert_round_open(db, invite.tenant_id, invite.round_id)


def _mark_in_progress(invite: AssessmentInvite) -> None:
    """HRP-358: the first saved score flips the invite to ``in_progress``.

    ``declined`` is terminal (HRP-359 REDO reversed the earlier change-of-
    mind allowance) — resolve_invite_by_token kills the token outright, so
    a declined invite can never reach a score write. No commit here — the
    caller's score write commits the session and the status change rides
    the same transaction.
    """
    if invite.status in ("pending", "opened"):
        invite.status = "in_progress"


async def public_get_context(
    db: AsyncSession, token: str, *, ip: str | None = None
) -> dict[str, Any]:
    from app.core.s3 import get_presigned_url
    from app.modules.auth.models import User
    from app.modules.company.models import Tenant
    from app.modules.recruitment.common import candidate_display_name
    from app.modules.recruitment.models import CandidateQuestion
    from app.modules.storage.models import File

    invite = await resolve_invite_by_token(db, token, ip=ip)
    cv = await db.get(CandidateVacancy, invite.candidate_vacancy_id)
    if cv is None:
        raise AppError("candidate_no_longer_available", status.HTTP_410_GONE)
    if cv.tenant_id != invite.tenant_id:
        raise AppError("cross_tenant_invite_mismatch", status.HTTP_403_FORBIDDEN)

    await _mark_opened(db, invite)

    tenant = await db.get(Tenant, invite.tenant_id)

    # Server-side consent gate (review [3/4]): candidate PII — name, the
    # presigned resume URL, interview questions — must not leave the API
    # until the evaluator has accepted the consent terms. Before that the
    # response carries only what the consent screen itself renders.
    if invite.consent_accepted_at is None:
        return {
            "invite_id": invite.id,
            "status": invite.status,
            "expires_at": invite.expires_at,
            "evaluator_name": invite.evaluator_name,
            "allow_reediting": invite.allow_reediting,
            "tenant_id": invite.tenant_id,
            "candidate_vacancy_id": invite.candidate_vacancy_id,
            "round_id": invite.round_id,
            "assessment_id": None,
            "personal_message": invite.personal_message,
            "consent_accepted": False,
            "tenant_name": tenant.name if tenant else None,
        }

    # Ensure the assessment sheet exists for this evaluator.
    assessment = None
    if invite.round_id is not None:
        assessment = await get_or_create_assessment(
            db,
            invite.tenant_id,
            invite.round_id,
            evaluator_invite_id=invite.id,
            evaluator_display_name=invite.evaluator_name or invite.email,
        )

    candidate = await db.get(Candidate, cv.candidate_id)
    vacancy = await db.get(Vacancy, cv.vacancy_id)

    # Latest resume file → short-lived presigned URL for the PDF pane.
    # ``file_type`` matters: the same table also holds interview media, and
    # the newest row on a candidate is often an audio recording — which
    # would then be handed to the evaluator as "the resume".
    resume_result = await db.execute(
        select(CandidateFile)
        .where(
            CandidateFile.candidate_id == cv.candidate_id,
            CandidateFile.tenant_id == invite.tenant_id,
            CandidateFile.file_type == "resume",
        )
        .order_by(CandidateFile.created_at.desc())
        .limit(1)
    )
    resume = resume_result.scalar_one_or_none()
    resume_url = None
    if resume and resume.file_id:
        file_record = await db.get(File, resume.file_id)
        if file_record:
            resume_url = get_presigned_url(file_record.path)

    q_result = await db.execute(
        select(CandidateQuestion)
        .where(
            CandidateQuestion.candidate_id == cv.candidate_id,
            CandidateQuestion.vacancy_id == cv.vacancy_id,
            CandidateQuestion.tenant_id == invite.tenant_id,
        )
        .order_by(CandidateQuestion.sort_order)
    )
    # Slimmed on purpose: the internal rubric (good/acceptable/poor answer
    # keys) is recruiter-only material and must not ship to an outsider.
    questions = [
        {
            "id": q.id,
            "question_text": q.question_text,
            "resume_fragment": q.resume_fragment,
            "purpose": q.purpose,
            "sort_order": q.sort_order,
        }
        for q in q_result.scalars().all()
    ]

    recruiter_name = None
    recruiter_email = None
    if invite.invited_by is not None:
        recruiter = await db.get(User, invite.invited_by)
        if recruiter is not None:
            recruiter_name = (
                f"{(recruiter.first_name or '').strip()} "
                f"{(recruiter.last_name or '').strip()}"
            ).strip() or None
            recruiter_email = recruiter.email

    # ``get_or_create_assessment`` snapshots the vacancy scale lazily, so
    # re-read the vacancy for a fresh snapshot before rendering the form.
    scale_levels: list[dict[str, Any]] = []
    if vacancy is not None and vacancy.assessment_scale_snapshot:
        scale_levels = vacancy.assessment_scale_snapshot.get("levels", [])

    competences = await list_vacancy_profile_competences(
        db, invite.tenant_id, cv.vacancy_id
    )

    return {
        "invite_id": invite.id,
        "status": invite.status,
        "expires_at": invite.expires_at,
        "evaluator_name": invite.evaluator_name,
        "allow_reediting": invite.allow_reediting,
        "tenant_id": invite.tenant_id,
        "candidate_vacancy_id": invite.candidate_vacancy_id,
        "round_id": invite.round_id,
        "assessment_id": assessment.id if assessment else None,
        "personal_message": invite.personal_message,
        "consent_accepted": invite.consent_accepted_at is not None,
        # HRP-359: everything the standalone evaluation page renders.
        "tenant_name": tenant.name if tenant else None,
        # i18n F7: no English "(unnamed)" on the wire — the public
        # evaluation page localizes the empty-name fallback itself.
        "candidate_name": candidate_display_name(candidate, fallback=""),
        "vacancy_title": vacancy.title if vacancy else None,
        "resume_url": resume_url,
        "resume_filename": resume.original_filename if resume else None,
        "resume_mime_type": resume.mime_type if resume else None,
        "questions": questions,
        "recruiter_name": recruiter_name,
        "recruiter_email": recruiter_email,
        "scale_levels": scale_levels,
        "competences": competences,
        "critical_submit_threshold": CRITICAL_SUBMIT_THRESHOLD,
        "assessment": (
            await get_assessment(db, invite.tenant_id, assessment.id)
            if assessment
            else None
        ),
    }


def _resume_kind(mime_type: str | None, filename: str | None) -> str:
    """Classify a resume file into a preview strategy.

    ``pdf`` renders inline in an iframe; ``text`` is extracted server-side;
    anything else only offers a download. The filename is a fallback for
    uploads that arrived without a usable mime type.
    """
    mime = (mime_type or "").lower()
    name = (filename or "").lower()
    if "pdf" in mime or name.endswith(".pdf"):
        return "pdf"
    if any(hint in mime for hint in _DOCX_MIME_HINTS) or name.endswith(".docx"):
        return "text"
    if any(hint in mime for hint in _PLAINTEXT_MIME_HINTS) or name.endswith(
        (".txt", ".rtf")
    ):
        return "text"
    return "unsupported"


def _docx_preview_blocks(data: bytes) -> list[str]:
    """Paragraph-level text of a DOCX, in reading order.

    Mirrors the resume-parsing extractor: many CV templates keep the whole
    document inside tables, so paragraphs alone would return an empty
    preview. Table rows collapse to ``cell | cell`` lines — enough to read
    a two-column resume without shipping a full HTML converter.
    """
    import io

    from docx import Document
    from docx.table import Table

    doc = Document(io.BytesIO(data))
    blocks: list[str] = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    def _walk_table(table: Table) -> None:
        # Merged cells surface once per spanned grid slot — dedupe on the
        # underlying XML element so the text is not repeated.
        seen: set[Any] = set()
        for row in table.rows:
            cells: list[str] = []
            for cell in row.cells:
                if cell._tc in seen:
                    continue
                seen.add(cell._tc)
                if cell.text.strip():
                    cells.append(cell.text.strip())
                for nested in cell.tables:
                    _walk_table(nested)
            if cells:
                blocks.append(" | ".join(cells))

    for table in doc.tables:
        _walk_table(table)

    for section in doc.sections:
        for para in (*section.header.paragraphs, *section.footer.paragraphs):
            if para.text.strip():
                blocks.append(para.text.strip())

    return blocks


def _plaintext_preview_blocks(data: bytes) -> list[str]:
    text = data.decode("utf-8", errors="ignore")
    return [line.strip() for line in text.splitlines() if line.strip()]


async def public_resume_preview(
    db: AsyncSession, token: str, *, ip: str | None = None
) -> dict[str, Any]:
    """Renderable preview of the candidate's resume for an invited evaluator.

    Read-only companion to :func:`public_get_context`, behind the same
    token *and* the same consent gate — the resume is candidate PII.
    """
    from app.core.s3 import download_bytes, get_presigned_url
    from app.modules.recruitment.candidate_service import _attachment_disposition
    from app.modules.storage.models import File

    invite = await resolve_invite_by_token(db, token, ip=ip)
    if invite.consent_accepted_at is None:
        raise AppError("assessment_consent_required", status.HTTP_403_FORBIDDEN)
    cv = await db.get(CandidateVacancy, invite.candidate_vacancy_id)
    if cv is None:
        raise AppError("candidate_no_longer_available", status.HTTP_410_GONE)
    if cv.tenant_id != invite.tenant_id:
        raise AppError("cross_tenant_invite_mismatch", status.HTTP_403_FORBIDDEN)

    empty: dict[str, Any] = {
        "kind": "none",
        "filename": None,
        "mime_type": None,
        "preview_url": None,
        "download_url": None,
        "blocks": [],
        "truncated": False,
    }

    resume_result = await db.execute(
        select(CandidateFile)
        .where(
            CandidateFile.candidate_id == cv.candidate_id,
            CandidateFile.tenant_id == invite.tenant_id,
            CandidateFile.file_type == "resume",
        )
        .order_by(CandidateFile.created_at.desc())
        .limit(1)
    )
    resume = resume_result.scalar_one_or_none()
    if resume is None or not resume.file_id:
        return empty
    file_record = await db.get(File, resume.file_id)
    if file_record is None:
        return empty

    kind = _resume_kind(resume.mime_type, resume.original_filename)
    # Content-Disposition: attachment is what makes "download" an explicit
    # user action rather than a side effect of opening the page. Same
    # RFC 5987 encoder the recruiter-facing download uses (HRP-347).
    download_url = get_presigned_url(
        file_record.path,
        content_disposition=_attachment_disposition(
            resume.original_filename or "resume"
        ),
    )
    out = {
        **empty,
        "kind": kind,
        "filename": resume.original_filename,
        "mime_type": resume.mime_type,
        "download_url": download_url,
    }
    if kind == "pdf":
        out["preview_url"] = get_presigned_url(file_record.path)
        return out
    if kind != "text":
        return out

    if (file_record.size or 0) > RESUME_PREVIEW_MAX_BYTES:
        # Refuse to pull an oversized document into the worker just to
        # render a preview; the evaluator can still download it.
        out["kind"] = "unsupported"
        return out
    # Both steps are synchronous and slow — up to 15 MB off S3 through
    # blocking boto3, then an XML walk over the whole document. Run on the
    # thread pool: this is an unauthenticated token endpoint, so leaving
    # them on the loop lets a handful of concurrent previews stall every
    # other request the worker is serving.
    data = await asyncio.to_thread(download_bytes, file_record.path)
    if not data:
        out["kind"] = "unsupported"
        return out
    try:
        mime = (resume.mime_type or "").lower()
        name = (resume.original_filename or "").lower()
        if any(hint in mime for hint in _DOCX_MIME_HINTS) or name.endswith(".docx"):
            blocks = await asyncio.to_thread(_docx_preview_blocks, data)
        else:
            blocks = await asyncio.to_thread(_plaintext_preview_blocks, data)
    except Exception:
        log.exception("resume preview extraction failed for invite=%s", invite.id)
        out["kind"] = "unsupported"
        return out

    if not blocks:
        # A readable file that yielded nothing (scanned pages inside a
        # docx, for instance) — an empty pane would look like a bug, so
        # fall back to the download-only affordance.
        out["kind"] = "unsupported"
        return out
    out["truncated"] = len(blocks) > RESUME_PREVIEW_MAX_BLOCKS
    out["blocks"] = blocks[:RESUME_PREVIEW_MAX_BLOCKS]
    return out


async def public_accept_consent(
    db: AsyncSession, token: str, *, ip: str | None = None
) -> dict[str, Any]:
    invite = await resolve_invite_by_token(db, token, ip=ip)
    if invite.status == "pending":
        invite.status = "opened"
    if invite.opened_at is None:
        invite.opened_at = datetime.now(timezone.utc)
    if invite.consent_accepted_at is None:
        invite.consent_accepted_at = datetime.now(timezone.utc)
    await db.commit()
    await db.refresh(invite)
    await audit_service.record_event(
        db,
        tenant_id=invite.tenant_id,
        user_id=None,
        action="assessment_invite.consent_accepted",
        entity_type="assessment_invite",
        entity_id=invite.id,
    )
    return {
        "status": invite.status,
        "consent_accepted": invite.consent_accepted_at is not None,
    }


async def public_decline(
    db: AsyncSession, token: str, *, ip: str | None = None
) -> dict[str, Any]:
    invite = await resolve_invite_by_token(db, token, ip=ip)
    if invite.status == "submitted":
        # A stale consent tab must not demote a completed evaluation whose
        # scores already feed the round aggregate (review [2]).
        raise AppError("evaluation_already_submitted", status.HTTP_409_CONFLICT)
    invite.status = "declined"
    invite.declined_at = datetime.now(timezone.utc)
    await db.commit()
    await db.refresh(invite)
    await audit_service.record_event(
        db,
        tenant_id=invite.tenant_id,
        user_id=None,
        action="assessment_invite.declined",
        entity_type="assessment_invite",
        entity_id=invite.id,
    )
    return {"status": invite.status}


async def public_update_name(
    db: AsyncSession, token: str, name: str, *, ip: str | None = None
) -> dict[str, Any]:
    invite = await resolve_invite_by_token(db, token, ip=ip)
    _ensure_editable(invite)
    await _assert_invite_round_open(db, invite)
    invite.evaluator_name = name
    await db.commit()
    if invite.round_id is not None:
        a = await get_or_create_assessment(
            db,
            invite.tenant_id,
            invite.round_id,
            evaluator_invite_id=invite.id,
            evaluator_display_name=name,
        )
        a.evaluator_display_name = name
        await db.commit()
    await audit_service.record_event(
        db,
        tenant_id=invite.tenant_id,
        user_id=None,
        action="assessment_invite.name_changed",
        entity_type="assessment_invite",
        entity_id=invite.id,
        payload_diff={"new_name_length": len(name)},
    )
    return {"evaluator_name": invite.evaluator_name}


async def public_submit(
    db: AsyncSession,
    token: str,
    *,
    ip: str | None = None,
    final_notes: str | None = None,
) -> dict[str, Any]:
    invite = await resolve_invite_by_token(db, token, ip=ip)
    _ensure_editable(invite)
    if invite.round_id is None:
        raise AppError("invite_has_no_round", status.HTTP_400_BAD_REQUEST)
    await _assert_invite_round_open(db, invite)
    a = await get_or_create_assessment(
        db,
        invite.tenant_id,
        invite.round_id,
        evaluator_invite_id=invite.id,
        evaluator_display_name=invite.evaluator_name or invite.email,
    )
    a.status = "submitted"
    a.submitted_at = datetime.now(timezone.utc)
    if final_notes is not None:
        a.final_notes = final_notes
    invite.status = "submitted"
    invite.submitted_at = datetime.now(timezone.utc)
    await db.commit()
    await db.refresh(invite)
    cv = await db.get(CandidateVacancy, invite.candidate_vacancy_id)
    if cv is not None:
        await recompute_manager_score(db, invite.tenant_id, cv.id)
    await audit_service.record_event(
        db,
        tenant_id=invite.tenant_id,
        user_id=None,
        action="assessment_invite.submitted",
        entity_type="assessment_invite",
        entity_id=invite.id,
    )
    return {"status": invite.status, "assessment_id": a.id}


async def public_save_final_notes(
    db: AsyncSession,
    token: str,
    final_notes: str,
    *,
    ip: str | None = None,
) -> dict[str, Any]:
    """Persist evaluator's free-form notes without changing assessment status.

    HRP-186 fix: the public form has an autosave indicator on the notes
    textarea; the previous implementation was a UI-only no-op, so notes
    were silently lost on refresh / accidental tab-close. Submit-time
    `final_notes` overwrite still applies — this endpoint just keeps the
    in-progress draft alive.
    """
    invite = await resolve_invite_by_token(db, token, ip=ip)
    _ensure_editable(invite)
    if invite.round_id is None:
        raise AppError("invite_has_no_round", status.HTTP_400_BAD_REQUEST)
    await _assert_invite_round_open(db, invite)
    a = await get_or_create_assessment(
        db,
        invite.tenant_id,
        invite.round_id,
        evaluator_invite_id=invite.id,
        evaluator_display_name=invite.evaluator_name or invite.email,
    )
    a.final_notes = final_notes
    await db.commit()
    return {"assessment_id": a.id, "final_notes": a.final_notes}


async def public_set_competence_score(
    db: AsyncSession,
    token: str,
    competence_id: uuid.UUID,
    score_value: int | None,
    comment: str | None,
    *,
    ip: str | None = None,
) -> dict[str, Any]:
    from app.modules.recruitment.manager_assessment_schemas import (
        CompetenceScoreIn,
    )
    from app.modules.recruitment.manager_assessment_service import (
        set_competence_score,
    )

    invite = await resolve_invite_by_token(db, token, ip=ip)
    _ensure_editable(invite)
    if invite.round_id is None:
        raise AppError("invite_has_no_round", status.HTTP_400_BAD_REQUEST)
    a = await get_or_create_assessment(
        db,
        invite.tenant_id,
        invite.round_id,
        evaluator_invite_id=invite.id,
        evaluator_display_name=invite.evaluator_name or invite.email,
    )
    if score_value is not None:
        _mark_in_progress(invite)
    return await set_competence_score(
        db,
        invite.tenant_id,
        None,
        a.id,
        competence_id,
        CompetenceScoreIn(
            score_value=score_value,
            score_source="manual",
            comment=comment,
        ),
    )


async def public_set_indicator_score(
    db: AsyncSession,
    token: str,
    indicator_id: uuid.UUID,
    competence_id: uuid.UUID,
    score_value: int | None,
    comment: str | None = None,
    *,
    ip: str | None = None,
) -> dict[str, Any]:
    """HRP-359: per-indicator scoring over the public token flow.

    Mirrors :func:`public_set_competence_score`; the service call also
    recomputes the competence overall from its indicators.
    """
    from app.modules.recruitment.manager_assessment_schemas import (
        IndicatorScoreIn,
    )
    from app.modules.recruitment.manager_assessment_service import (
        set_indicator_score,
    )

    invite = await resolve_invite_by_token(db, token, ip=ip)
    _ensure_editable(invite)
    if invite.round_id is None:
        raise AppError("invite_has_no_round", status.HTTP_400_BAD_REQUEST)
    a = await get_or_create_assessment(
        db,
        invite.tenant_id,
        invite.round_id,
        evaluator_invite_id=invite.id,
        evaluator_display_name=invite.evaluator_name or invite.email,
    )
    if score_value is not None:
        _mark_in_progress(invite)
    return await set_indicator_score(
        db,
        invite.tenant_id,
        None,
        a.id,
        indicator_id,
        IndicatorScoreIn(
            competence_id=competence_id,
            score_value=score_value,
            comment=comment,
        ),
    )


__all__ = [
    "resolve_invite_by_token",
    "public_get_context",
    "public_resume_preview",
    "public_accept_consent",
    "public_decline",
    "public_update_name",
    "public_save_final_notes",
    "public_submit",
    "public_set_competence_score",
    "public_set_indicator_score",
    "is_ip_blocked",
    "register_invalid_attempt",
    "PUBLIC_RATE_LIMIT_INVALID_THRESHOLD",
    "PUBLIC_RATE_LIMIT_WINDOW_MINUTES",
    "PUBLIC_RATE_LIMIT_BLOCK_HOURS",
    "CRITICAL_SUBMIT_THRESHOLD",
    "RESUME_PREVIEW_MAX_BLOCKS",
]
