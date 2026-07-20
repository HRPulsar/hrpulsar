"""Shared file parser for AI generation flows.

Whitelist: ``.pdf .docx .xlsx .xls .rtf .txt``. ``.doc`` (Word 97-2003 binary)
is *not* supported — there is no maintained Python library for it and the
LibreOffice subprocess approach pulls a 400 MB dependency. The router
returns 400 and asks the user to re-save as ``.docx`` or ``.pdf``.

Limits enforced here:

* per-file: 10 MiB (request body cap is 1 file, but we re-check after
  the framework has buffered it);
* aggregate extracted text: 50 000 characters across all files in one call.

Each extractor is best-effort plain-text only: tables are flattened into
space-separated rows, formatting is stripped. The downstream LLM gets a
single large string per file plus a header naming the file. Anything that
fails to parse raises ``UnsupportedFileError`` (400) — never silently
returns empty text, otherwise we'd burn LLM credits on garbage input.
"""

from __future__ import annotations

import io
import logging
import re
from typing import TYPE_CHECKING

logger = logging.getLogger(__name__)

if TYPE_CHECKING:  # pragma: no cover
    pass

SUPPORTED_EXTENSIONS: tuple[str, ...] = (
    ".pdf",
    ".docx",
    ".xlsx",
    ".xls",
    ".rtf",
    ".txt",
)
MAX_FILE_BYTES = 10 * 1024 * 1024
MAX_TOTAL_CHARS = 50_000

UNSUPPORTED_MESSAGE = (
    "Supported file types: .pdf, .docx, .xlsx, .xls, .rtf, .txt. "
    "Save the file in one of these formats."
)


class UnsupportedFileError(ValueError):
    """Whitelist or size-limit violation.

    Callers MUST catch this and translate to HTTP 400 — FastAPI does not
    auto-map ``ValueError`` subclasses to 4xx, so leaving it uncaught
    surfaces as a generic 500. See ``router.create_matrix_session`` for
    the canonical try/except pattern.
    """


class FileParseError(RuntimeError):
    """Whitelisted file failed to parse (corrupt content / empty PDF).

    Same caller contract as ``UnsupportedFileError``: translate to HTTP
    400 explicitly, otherwise FastAPI returns 500.
    """


def _ext(filename: str) -> str:
    name = (filename or "").strip().lower()
    if "." not in name:
        return ""
    return "." + name.rsplit(".", 1)[1]


def _parse_txt(content: bytes) -> str:
    # BOM-prefixed UTF-16 first (windows-style "Save as Unicode"), then UTF-8,
    # finally latin-1 as a permissive fallback for legacy Western European
    # exports. UTF-16 without BOM decodes almost any byte stream, so we only
    # take it when a BOM signals it explicitly.
    if content.startswith((b"\xff\xfe", b"\xfe\xff")):
        try:
            return content.decode("utf-16")
        except UnicodeDecodeError:
            pass
    for encoding in ("utf-8", "latin-1"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="replace")


_RTF_CONTROL = re.compile(r"\\[a-zA-Z]+-?\d* ?|\\\*|\\['\"]")
_RTF_GROUPS = re.compile(r"[{}]")


def _parse_rtf(content: bytes) -> str:
    """Strip RTF control words and braces. Best-effort, not a full parser."""
    text = _parse_txt(content)
    text = _RTF_CONTROL.sub("", text)
    text = _RTF_GROUPS.sub("", text)
    return re.sub(r"\s+", " ", text).strip()


def _parse_pdf(content: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover — runtime dependency
        raise FileParseError("pypdf is not installed") from exc

    try:
        reader = PdfReader(io.BytesIO(content))
    except Exception as exc:
        raise FileParseError(f"Failed to parse PDF: {exc}") from exc

    pages: list[str] = []
    failed_pages = 0
    total_pages = 0
    for page in reader.pages:
        total_pages += 1
        try:
            pages.append(page.extract_text() or "")
        except Exception:
            failed_pages += 1
            logger.exception("pdf page extraction failed")

    text = "\n".join(p for p in pages if p)
    if total_pages > 0 and not text:
        # Per-page exceptions are tolerated, but if every page produced
        # nothing the file is unusable — refuse upstream so we don't burn
        # LLM credits on an empty parsed_files block.
        raise FileParseError(
            f"Failed to extract any text from PDF "
            f"({failed_pages}/{total_pages} pages errored)."
        )
    return text


def _parse_docx(content: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover
        raise FileParseError("python-docx is not installed") from exc

    try:
        doc = Document(io.BytesIO(content))
        parts = [para.text for para in doc.paragraphs if para.text]
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    except Exception as exc:
        raise FileParseError(f"Failed to parse DOCX: {exc}") from exc


def _parse_xlsx(content: bytes) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:  # pragma: no cover
        raise FileParseError("openpyxl is not installed") from exc

    try:
        wb = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
        parts: list[str] = []
        for sheet in wb.worksheets:
            parts.append(f"# Sheet: {sheet.title}")
            for row in sheet.iter_rows(values_only=True):
                cells = [str(v) for v in row if v is not None and str(v).strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    except Exception as exc:
        raise FileParseError(f"Failed to parse XLSX: {exc}") from exc


def _parse_xls(content: bytes) -> str:
    try:
        import xlrd
    except ImportError as exc:  # pragma: no cover
        raise FileParseError("xlrd is not installed") from exc

    try:
        wb = xlrd.open_workbook(file_contents=content)
        parts: list[str] = []
        for sheet in wb.sheets():
            parts.append(f"# Sheet: {sheet.name}")
            for row_idx in range(sheet.nrows):
                row = sheet.row_values(row_idx)
                cells = [str(v) for v in row if v not in (None, "") and str(v).strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    except Exception as exc:
        raise FileParseError(f"Failed to parse XLS: {exc}") from exc


_PARSERS = {
    ".txt": _parse_txt,
    ".rtf": _parse_rtf,
    ".pdf": _parse_pdf,
    ".docx": _parse_docx,
    ".xlsx": _parse_xlsx,
    ".xls": _parse_xls,
}


def parse_file(filename: str, content: bytes) -> str:
    """Parse one file. Raises UnsupportedFileError on whitelist/size violations."""
    ext = _ext(filename)
    if ext not in SUPPORTED_EXTENSIONS:
        raise UnsupportedFileError(UNSUPPORTED_MESSAGE)
    if len(content) > MAX_FILE_BYTES:
        size_mb = len(content) / (1024 * 1024)
        raise UnsupportedFileError(
            f"File '{filename}' is {size_mb:.1f} MB — limit is "
            f"{MAX_FILE_BYTES // (1024 * 1024)} MB."
        )
    parser = _PARSERS[ext]
    text = parser(content)
    if isinstance(text, bytes):  # pragma: no cover — defensive
        text = text.decode("utf-8", errors="replace")
    return text.strip()


def parse_files(files: list[tuple[str, bytes]]) -> str:
    """Parse a batch and concatenate results, headed by filename.

    Aggregate text length is capped at MAX_TOTAL_CHARS — anything beyond
    raises UnsupportedFileError so the router returns 400 *before* we
    spend credits on an LLM call we'd have to truncate anyway.
    """
    chunks: list[str] = []
    total = 0
    for filename, content in files:
        text = parse_file(filename, content)
        if not text:
            continue
        block = f"# File: {filename}\n{text}"
        total += len(block)
        if total > MAX_TOTAL_CHARS:
            raise UnsupportedFileError(
                f"Combined file content exceeds {MAX_TOTAL_CHARS} characters. "
                "Trim the inputs and retry."
            )
        chunks.append(block)
    return "\n\n".join(chunks)
