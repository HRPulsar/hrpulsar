"""Unit tests for `app.modules.ai.file_parsing` (POS5).

Whitelist + size limits + per-format parsing (txt/rtf/docx/xlsx — pdf and
xls are skipped here because building a fixture for them requires the
upstream libraries; the parser dispatch is covered by the explicit ext
mapping unit test).
"""

from __future__ import annotations

import io
import zipfile

import pytest
from app.modules.ai import file_parsing
from openpyxl import Workbook


def test_supported_extensions_whitelist():
    assert set(file_parsing.SUPPORTED_EXTENSIONS) == {
        ".pdf",
        ".docx",
        ".xlsx",
        ".xls",
        ".rtf",
        ".txt",
    }


class TestParseFile:
    def test_txt_basic(self):
        text = file_parsing.parse_file("notes.txt", b"Hello world\nLine two")
        assert "Hello world" in text
        assert "Line two" in text

    def test_txt_latin1_decoded(self):
        # "Café" encoded as latin-1 — the 0xE9 byte for "é" is not a valid
        # UTF-8 start byte, so the utf-8 decoder fails and the latin-1
        # fallback wins.
        content = "Café".encode("latin-1")
        assert b"\xe9" in content  # sanity: confirms it's not UTF-8
        text = file_parsing.parse_file("notes.txt", content)
        assert "Café" in text

    def test_rtf_strips_control_words(self):
        rtf = (
            rb"{\rtf1\ansi\deff0 {\fonttbl {\f0 Times New Roman;}}"
            rb"\f0\fs24 Hello \b world\b0 .}"
        )
        text = file_parsing.parse_file("doc.rtf", rtf)
        assert "Hello" in text
        assert "world" in text
        assert "rtf1" not in text

    def test_doc_extension_rejected(self):
        with pytest.raises(file_parsing.UnsupportedFileError) as exc:
            file_parsing.parse_file("legacy.doc", b"\xd0\xcf\x11\xe0...")
        msg = str(exc.value)
        assert ".pdf" in msg and ".docx" in msg

    def test_unknown_extension_rejected(self):
        with pytest.raises(file_parsing.UnsupportedFileError):
            file_parsing.parse_file("malware.exe", b"MZ")

    def test_no_extension_rejected(self):
        with pytest.raises(file_parsing.UnsupportedFileError):
            file_parsing.parse_file("README", b"Hello")

    def test_oversized_file_rejected(self):
        big = b"a" * (file_parsing.MAX_FILE_BYTES + 1)
        with pytest.raises(file_parsing.UnsupportedFileError) as exc:
            file_parsing.parse_file("big.txt", big)
        assert "limit" in str(exc.value).lower()

    def test_xlsx_extracts_cells(self):
        wb = Workbook()
        ws = wb.active
        ws.title = "Skills"
        ws.append(["Competence", "Level"])
        ws.append(["API design", "intermediate"])
        ws.append(["SQL", "basic"])
        buf = io.BytesIO()
        wb.save(buf)

        text = file_parsing.parse_file("matrix.xlsx", buf.getvalue())
        assert "Skills" in text
        assert "API design" in text
        assert "intermediate" in text

    def test_docx_extracts_paragraphs(self):
        from docx import Document

        doc = Document()
        doc.add_paragraph("Backend engineer responsibilities.")
        doc.add_paragraph("Owns API layer.")
        buf = io.BytesIO()
        doc.save(buf)

        text = file_parsing.parse_file("role.docx", buf.getvalue())
        assert "Backend engineer responsibilities" in text
        assert "Owns API layer" in text

    def test_pdf_with_no_extractable_text_raises(self, monkeypatch):
        """Per-page exceptions are tolerated; full failure must raise."""
        import app.modules.ai.file_parsing as fp

        class _BadPage:
            def extract_text(self):
                raise RuntimeError("page broken")

        class _StubReader:
            def __init__(self, *_a, **_kw):
                self.pages = [_BadPage(), _BadPage()]

        # `_parse_pdf` does `from pypdf import PdfReader` inside the function;
        # patch that name on the pypdf module so the inline import resolves
        # to our stub.
        import pypdf

        monkeypatch.setattr(pypdf, "PdfReader", _StubReader)

        with pytest.raises(fp.FileParseError):
            fp.parse_file("broken.pdf", b"%PDF-1.4\n...stub...")

    def test_corrupt_docx_raises_parse_error(self):
        # A zip file that isn't a docx — Document() construction fails.
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w") as zf:
            zf.writestr("not-a-docx.txt", b"random")
        with pytest.raises(file_parsing.FileParseError):
            file_parsing.parse_file("fake.docx", buf.getvalue())


class TestParseFiles:
    def test_aggregate_under_limit(self):
        files = [
            ("a.txt", b"alpha alpha"),
            ("b.txt", b"bravo bravo"),
        ]
        out = file_parsing.parse_files(files)
        assert "# File: a.txt" in out
        assert "# File: b.txt" in out
        assert "alpha" in out and "bravo" in out

    def test_aggregate_overflow_rejected(self):
        big = b"x" * (file_parsing.MAX_TOTAL_CHARS + 100)
        files = [("big.txt", big)]
        with pytest.raises(file_parsing.UnsupportedFileError) as exc:
            file_parsing.parse_files(files)
        assert "characters" in str(exc.value)

    def test_doc_inside_batch_rejected_early(self):
        files = [("ok.txt", b"hello"), ("legacy.doc", b"binary")]
        with pytest.raises(file_parsing.UnsupportedFileError):
            file_parsing.parse_files(files)

    def test_empty_batch(self):
        assert file_parsing.parse_files([]) == ""
