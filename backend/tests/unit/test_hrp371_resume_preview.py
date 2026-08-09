"""HRP-371: the invited evaluator gets a resume *preview*, not a download.

Pointing the pane's iframe at the presigned URL of a .docx left the block
empty and dropped the file into the evaluator's Downloads folder on every
page load. The server now classifies the file and extracts text for the
formats a browser cannot render.
"""

from __future__ import annotations

import io
import uuid
from datetime import datetime, timedelta, timezone
from unittest.mock import patch

import pytest
from app.modules.recruitment import manager_assessment_public as public_service
from app.modules.recruitment.models import (
    AssessmentInvite,
    CandidateFile,
)
from app.modules.storage.models import File
from fastapi import HTTPException
from sqlalchemy.ext.asyncio import AsyncSession

from tests.unit.test_hrp186_manager_assessment import (
    _make_candidate_vacancy,
    _make_profile,
    _make_vacancy,
)

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _docx_bytes(paragraphs: list[str], table_rows: list[list[str]] | None = None):
    """A real .docx built with the same library the extractor reads."""
    from docx import Document

    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    if table_rows:
        table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for r, row in enumerate(table_rows):
            for c, value in enumerate(row):
                table.cell(r, c).text = value
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


async def _invited_evaluator(
    db: AsyncSession,
    tenant,
    user,
    *,
    consent: bool = True,
    filename: str = "cv.docx",
    mime_type: str = DOCX_MIME,
    size: int = 4096,
    with_file: bool = True,
):
    """A live invite on a candidate that carries one resume file."""
    vacancy = await _make_vacancy(db, tenant)
    await _make_profile(db, tenant, vacancy)
    cv = await _make_candidate_vacancy(db, tenant, vacancy)

    if with_file:
        file_row = File(
            tenant_id=tenant.id,
            name=filename,
            original_name=filename,
            path=f"{tenant.id}/resumes/{uuid.uuid4().hex}",
            size=size,
            mime_type=mime_type,
            uploaded_by=user.id,
        )
        db.add(file_row)
        await db.flush()
        db.add(
            CandidateFile(
                tenant_id=tenant.id,
                candidate_id=cv.candidate_id,
                file_id=file_row.id,
                file_type="resume",
                original_filename=filename,
                mime_type=mime_type,
                file_size=size,
            )
        )

    token = uuid.uuid4().hex
    invite = AssessmentInvite(
        tenant_id=tenant.id,
        candidate_vacancy_id=cv.id,
        token=token,
        token_hash=public_service.hash_token(token),
        email="ext@example.com",
        evaluator_name="Ext Eval",
        status="opened",
        expires_at=datetime.now(timezone.utc) + timedelta(days=7),
        allow_reediting=True,
        delivery_status="sent",
        consent_accepted_at=datetime.now(timezone.utc) if consent else None,
    )
    db.add(invite)
    await db.commit()
    return token, cv


class TestResumeKindClassification:
    def test_pdf_by_mime_and_by_extension(self):
        assert public_service._resume_kind("application/pdf", "x.bin") == "pdf"
        assert public_service._resume_kind(None, "Resume.PDF") == "pdf"

    def test_docx_is_text_extractable(self):
        assert public_service._resume_kind(DOCX_MIME, "cv.docx") == "text"
        assert public_service._resume_kind(None, "cv.docx") == "text"

    def test_plaintext_and_rtf(self):
        assert public_service._resume_kind("text/plain", "cv.txt") == "text"
        assert public_service._resume_kind("text/rtf", "cv.rtf") == "text"

    def test_legacy_doc_is_unsupported(self):
        assert public_service._resume_kind("application/msword", "cv.doc") == (
            "unsupported"
        )


class TestDocxExtraction:
    def test_paragraphs_are_returned_in_order(self):
        blocks = public_service._docx_preview_blocks(
            _docx_bytes(["Jane Doe", "Senior Engineer", "  "])
        )
        assert blocks[:2] == ["Jane Doe", "Senior Engineer"]
        # Blank paragraphs are dropped rather than rendered as empty lines.
        assert "" not in blocks

    def test_table_only_resume_is_not_empty(self):
        # Two-column CV templates keep everything in a table; paragraphs
        # alone would return nothing and the pane would look broken.
        blocks = public_service._docx_preview_blocks(
            _docx_bytes([], table_rows=[["Skills", "Python, SQL"]])
        )
        assert any("Python, SQL" in b for b in blocks)


class TestResumePreviewEndpoint:
    async def test_docx_returns_text_blocks_and_no_inline_url(
        self, db: AsyncSession, tenant, user
    ):
        token, _ = await _invited_evaluator(db, tenant, user)
        data = _docx_bytes(["Jane Doe", "10 years of Python"])
        with (
            patch("app.core.s3.download_bytes", return_value=data),
            patch("app.core.s3.get_presigned_url", return_value="https://s3/signed"),
        ):
            out = await public_service.public_resume_preview(db, token)

        assert out["kind"] == "text"
        assert "Jane Doe" in out["blocks"]
        # The pane must not be handed a URL it would drop into an iframe —
        # that is exactly what triggered the silent download.
        assert out["preview_url"] is None
        assert out["download_url"] == "https://s3/signed"
        assert out["truncated"] is False

    async def test_download_url_forces_attachment_disposition(
        self, db: AsyncSession, tenant, user
    ):
        token, _ = await _invited_evaluator(db, tenant, user)
        with (
            patch("app.core.s3.download_bytes", return_value=_docx_bytes(["x"])),
            patch(
                "app.core.s3.get_presigned_url", return_value="https://s3/signed"
            ) as presign,
        ):
            await public_service.public_resume_preview(db, token)

        disposition = presign.call_args_list[0].kwargs["content_disposition"]
        assert disposition.startswith("attachment;")
        assert "cv.docx" in disposition

    async def test_pdf_keeps_the_inline_iframe_url(
        self, db: AsyncSession, tenant, user
    ):
        token, _ = await _invited_evaluator(
            db, tenant, user, filename="cv.pdf", mime_type="application/pdf"
        )
        with patch("app.core.s3.get_presigned_url", return_value="https://s3/signed"):
            out = await public_service.public_resume_preview(db, token)

        assert out["kind"] == "pdf"
        assert out["preview_url"] == "https://s3/signed"
        assert out["blocks"] == []

    async def test_unreadable_docx_degrades_to_download_only(
        self, db: AsyncSession, tenant, user
    ):
        token, _ = await _invited_evaluator(db, tenant, user)
        with (
            patch("app.core.s3.download_bytes", return_value=b"not-a-docx"),
            patch("app.core.s3.get_presigned_url", return_value="https://s3/signed"),
        ):
            out = await public_service.public_resume_preview(db, token)

        assert out["kind"] == "unsupported"
        assert out["download_url"] == "https://s3/signed"

    async def test_oversized_file_is_not_pulled_into_the_worker(
        self, db: AsyncSession, tenant, user
    ):
        token, _ = await _invited_evaluator(
            db, tenant, user, size=public_service.RESUME_PREVIEW_MAX_BYTES + 1
        )
        with (
            patch("app.core.s3.download_bytes") as download,
            patch("app.core.s3.get_presigned_url", return_value="https://s3/signed"),
        ):
            out = await public_service.public_resume_preview(db, token)

        assert out["kind"] == "unsupported"
        download.assert_not_called()

    async def test_candidate_without_resume(self, db: AsyncSession, tenant, user):
        token, _ = await _invited_evaluator(db, tenant, user, with_file=False)
        out = await public_service.public_resume_preview(db, token)
        assert out["kind"] == "none"
        assert out["download_url"] is None

    async def test_consent_gate_blocks_the_resume(self, db: AsyncSession, tenant, user):
        token, _ = await _invited_evaluator(db, tenant, user, consent=False)
        with pytest.raises(HTTPException) as exc:
            await public_service.public_resume_preview(db, token)
        assert exc.value.status_code == 403

    async def test_unknown_token_is_gone(self, db: AsyncSession, tenant):
        with pytest.raises(HTTPException) as exc:
            await public_service.public_resume_preview(db, "nope")
        assert exc.value.status_code == 410

    async def test_revoked_token_is_gone(self, db: AsyncSession, tenant, user):
        from sqlalchemy import select

        token, _ = await _invited_evaluator(db, tenant, user)
        invite = (
            await db.execute(
                select(AssessmentInvite).where(
                    AssessmentInvite.token_hash == public_service.hash_token(token)
                )
            )
        ).scalar_one()
        invite.status = "revoked"
        invite.revoked_at = datetime.now(timezone.utc)
        await db.commit()

        with pytest.raises(HTTPException) as exc:
            await public_service.public_resume_preview(db, token)
        assert exc.value.status_code == 410

    async def test_long_document_is_truncated(self, db: AsyncSession, tenant, user):
        token, _ = await _invited_evaluator(db, tenant, user)
        paragraphs = [
            f"line {i}" for i in range(public_service.RESUME_PREVIEW_MAX_BLOCKS + 25)
        ]
        with (
            patch("app.core.s3.download_bytes", return_value=_docx_bytes(paragraphs)),
            patch("app.core.s3.get_presigned_url", return_value="https://s3/signed"),
        ):
            out = await public_service.public_resume_preview(db, token)

        assert out["truncated"] is True
        assert len(out["blocks"]) == public_service.RESUME_PREVIEW_MAX_BLOCKS
