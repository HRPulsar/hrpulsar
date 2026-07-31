"""HRP-181 REDO Stage 3 — bulk resume upload + batch LLM parsing.

Covers the upload limits, the detached parsing flow, the legacy single-
file flow's back-compat path, parsing-status / dedup-preview reads, and
the denormalisation of LLM-extracted fields into the canonical Candidate
when ``from-parsed`` finalises a batch.
"""

from __future__ import annotations

import asyncio
import uuid
from datetime import datetime, timedelta, timezone
from io import BytesIO

import pytest
from app.modules.recruitment import ai_service, service, tasks
from app.modules.recruitment.models import (
    Candidate,
    CandidateFile,
    CandidateVacancy,
)
from app.modules.recruitment.prompts import PARSE_RESUME
from app.modules.recruitment.schemas import (
    BatchFinalizeFile,
    BatchFinalizeRequest,
    VacancyCreate,
)
from fastapi import HTTPException, UploadFile
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from starlette.datastructures import Headers

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


async def _make_vacancy(db: AsyncSession, tenant_id, user_id):
    await service.seed_default_recruitment_stages(db, tenant_id)
    await db.commit()
    return await service.create_vacancy(
        db,
        tenant_id,
        user_id,
        VacancyCreate(title=f"V-{uuid.uuid4().hex[:6]}"),
    )


async def _run_parse_task(*args) -> dict:
    """Invoke the Celery task body in a worker thread so the ``asyncio.run``
    inside the task can spin up its own event loop without colliding with
    pytest-asyncio's running loop."""

    return await asyncio.to_thread(tasks.parse_resume_task.run, *args)


def _make_upload(*, filename: str, content_type: str, body: bytes) -> UploadFile:
    file = UploadFile(
        file=BytesIO(body),
        filename=filename,
        headers=Headers({"content-type": content_type}),
    )
    # The bulk endpoint reads ``upload.size`` for the 25 MB pre-check and
    # again for billing — Starlette only fills it on real ASGI requests,
    # so we set it manually here.
    file.size = len(body)
    return file


def _pdf(body_extra: bytes = b"") -> bytes:
    return b"%PDF-1.4\n" + body_extra + b"%%EOF"


def _docx(body_extra: bytes = b"") -> bytes:
    return b"PK\x03\x04" + body_extra + b"docx-stub"


@pytest.fixture(autouse=True)
def _stub_s3_and_celery(monkeypatch):
    """Keep tests away from real S3 / Celery. The service imports both
    lazily inside ``bulk_upload_resumes`` so the patch must point at the
    full module path the call resolves at runtime. Also rewire the sync
    DB engine the Celery task spins up to the test database, so it can
    see rows committed via the async test session (mirrors the pattern
    in test_recruitment_reports.py)."""
    from app.config import settings as app_settings

    from tests.conftest import TEST_DB_URL

    monkeypatch.setattr(app_settings, "database_url", TEST_DB_URL)
    monkeypatch.setattr(
        "app.core.s3.upload_file",
        lambda *args, **kwargs: True,
    )
    monkeypatch.setattr(
        "app.modules.recruitment.tasks.parse_resume_task.delay",
        lambda *args, **kwargs: None,
    )


# ---------------------------------------------------------------------------
# Bulk upload limits + happy path
# ---------------------------------------------------------------------------


class TestBulkUpload:
    async def test_happy_path_creates_detached_rows(
        self, db: AsyncSession, tenant, user, monkeypatch
    ):
        vacancy = await _make_vacancy(db, tenant.id, user.id)
        calls: list[tuple[str, str]] = []
        monkeypatch.setattr(
            "app.modules.recruitment.tasks.parse_resume_task.delay",
            lambda file_id, tenant_id: calls.append((file_id, tenant_id)),
        )

        files = [
            _make_upload(
                filename=f"r{i}.pdf",
                content_type="application/pdf",
                body=_pdf(b"x" * 100),
            )
            for i in range(3)
        ]
        result = await service.bulk_upload_resumes(
            db, tenant.id, user.id, vacancy["id"], files
        )

        assert len(result) == 3
        for ack in result:
            assert ack["parse_status"] == "pending"
            assert ack["mime_type"] == "application/pdf"
            assert ack["file_size"] > 0

        rows = (
            (
                await db.execute(
                    select(CandidateFile).where(
                        CandidateFile.tenant_id == tenant.id,
                        CandidateFile.candidate_id.is_(None),
                    )
                )
            )
            .scalars()
            .all()
        )
        assert len(rows) == 3
        assert all(r.file_type == "resume" for r in rows)
        # B6 restoration: invariants previously covered by the legacy
        # TestResumeUpload class. Each freshly uploaded row enters DB with
        # ``parsed_data`` empty and ``parse_status='pending'`` so the
        # poller doesn't prematurely flip to ``ready``.
        assert all(r.parsed_data is None for r in rows)
        assert all(r.parse_status == "pending" for r in rows)
        # Each detached file enqueued exactly one Celery task.
        assert len(calls) == 3
        assert {tenant_id for _, tenant_id in calls} == {str(tenant.id)}

    async def test_rejects_over_limit_files(self, db: AsyncSession, tenant, user):
        vacancy = await _make_vacancy(db, tenant.id, user.id)
        files = [
            _make_upload(
                filename=f"r{i}.pdf",
                content_type="application/pdf",
                body=_pdf(),
            )
            for i in range(service.MAX_BULK_RESUME_FILES + 1)
        ]
        with pytest.raises(HTTPException) as exc:
            await service.bulk_upload_resumes(
                db, tenant.id, user.id, vacancy["id"], files
            )
        assert exc.value.status_code == 413

    async def test_rejects_oversized_file(self, db: AsyncSession, tenant, user):
        vacancy = await _make_vacancy(db, tenant.id, user.id)
        big = _pdf(b"\x00" * (service.MAX_RESUME_BYTES + 1))
        upload = _make_upload(
            filename="big.pdf", content_type="application/pdf", body=big
        )
        with pytest.raises(HTTPException) as exc:
            await service.bulk_upload_resumes(
                db, tenant.id, user.id, vacancy["id"], [upload]
            )
        assert exc.value.status_code == 413

    async def test_rejects_batch_total_above_cap(self, db: AsyncSession, tenant, user):
        """The aggregate batch cap fires once individual files are valid."""
        vacancy = await _make_vacancy(db, tenant.id, user.id)
        # Each PDF lands just under the per-file cap; 11 × ~10 MB =
        # ~110 MB, which trips the 100 MB total guard.
        chunk = b"\x00" * (service.MAX_RESUME_BYTES - 1024)
        files = [
            _make_upload(
                filename=f"r{i}.pdf",
                content_type="application/pdf",
                body=_pdf(chunk),
            )
            for i in range(11)
        ]
        with pytest.raises(HTTPException) as exc:
            await service.bulk_upload_resumes(
                db, tenant.id, user.id, vacancy["id"], files
            )
        assert exc.value.status_code == 413
        assert "total limit" in exc.value.detail.lower()

    async def test_rejects_unsupported_mime(self, db: AsyncSession, tenant, user):
        vacancy = await _make_vacancy(db, tenant.id, user.id)
        upload = _make_upload(
            filename="cv.txt",
            content_type="text/plain",
            body=b"hello",
        )
        with pytest.raises(HTTPException) as exc:
            await service.bulk_upload_resumes(
                db, tenant.id, user.id, vacancy["id"], [upload]
            )
        assert exc.value.status_code == 415

    async def test_rejects_mime_spoof(self, db: AsyncSession, tenant, user):
        """PDF Content-Type but the bytes are not a PDF — sniff catches it."""
        vacancy = await _make_vacancy(db, tenant.id, user.id)
        upload = _make_upload(
            filename="fake.pdf",
            content_type="application/pdf",
            body=b"this is plain text, no PDF magic header",
        )
        with pytest.raises(HTTPException) as exc:
            await service.bulk_upload_resumes(
                db, tenant.id, user.id, vacancy["id"], [upload]
            )
        assert exc.value.status_code == 415


# ---------------------------------------------------------------------------
# Celery task — detached vs. legacy
# ---------------------------------------------------------------------------


class _SyncFile:
    """Stand-in for ``CandidateFile`` written by an asyncpg test session
    but read by the sync Celery worker session. The sync engine uses the
    same Postgres database, so ``parse_resume_task`` happily picks the
    row up by id after we commit on the async side."""


class TestParseResumeTask:
    def _install_retry_stub(self, monkeypatch, retries: list[Exception]):
        """Replace ``Task.retry`` with a stub that re-raises so the calling
        test can detect retry behaviour without bringing Celery up.

        ``parse_resume_task.run`` calls ``self.retry`` on failure; we want
        to surface the exception as a normal Python error rather than a
        ``RetryTaskError`` that needs a worker to interpret.
        """

        def _fake_retry(exc=None, **kwargs):  # noqa: D401
            if exc is not None:
                retries.append(exc)
                raise exc
            raise RuntimeError("retry without exc")

        # ``parse_resume_task.retry`` lives on the task instance, so the
        # patched attribute is looked up as an instance method without
        # ``self`` binding — keep the stub's signature parameter-only.
        monkeypatch.setattr(tasks.parse_resume_task, "retry", _fake_retry)

    async def test_detached_file_no_candidate_touch(
        self, db: AsyncSession, tenant, user, monkeypatch
    ):
        """Detached file → parsed_data lands, no Candidate / Person work."""

        vacancy = await _make_vacancy(db, tenant.id, user.id)
        upload = _make_upload(
            filename="cv.pdf", content_type="application/pdf", body=_pdf()
        )
        result = await service.bulk_upload_resumes(
            db, tenant.id, user.id, vacancy["id"], [upload]
        )
        await db.commit()
        file_id = result[0]["file_id"]

        # Stub LLM + text extraction.
        monkeypatch.setattr(
            "app.modules.recruitment.tasks.parsing._extract_text_from_s3",
            lambda *args, **kwargs: "some resume text",
        )

        async def fake_parse(text: str, **kwargs) -> dict:
            return {
                "first_name": "Maya",
                "last_name": "K",
                "contacts": {"email": "maya@example.com"},
                "experience": [{"position": "Lead", "role": "Lead"}],
                "current_position": "Lead",
                "years_of_experience": 9,
                "location": "Berlin",
            }

        monkeypatch.setattr(
            "app.modules.recruitment.ai_service.parse_resume_text", fake_parse
        )

        # Invoke task body directly (skip Celery dispatch).
        retries: list[Exception] = []
        self._install_retry_stub(monkeypatch, retries)
        outcome = await _run_parse_task(str(file_id), str(tenant.id))
        assert outcome["status"] == "completed"
        assert retries == []

        # Refresh from the async session.
        await db.commit()
        row = (
            await db.execute(select(CandidateFile).where(CandidateFile.id == file_id))
        ).scalar_one()
        assert row.parse_status == "completed"
        assert row.candidate_id is None
        assert row.parsed_data["current_position"] == "Lead"
        # No Candidate row was created for the detached parse.
        candidates = (
            (
                await db.execute(
                    select(Candidate).where(Candidate.tenant_id == tenant.id)
                )
            )
            .scalars()
            .all()
        )
        assert candidates == []

    async def test_legacy_candidate_backfills_only_if_empty(
        self, db: AsyncSession, tenant, user, monkeypatch
    ):
        """Legacy candidate_id set → parsed_resume_jsonb backfills when None.
        Person is never touched."""

        existing = Candidate(
            tenant_id=tenant.id,
            full_name="Has Parsed",
            email="hp@example.com",
            parsed_resume_jsonb={"existing": True},
        )
        empty = Candidate(
            tenant_id=tenant.id,
            full_name="No Parsed",
            email="np@example.com",
        )
        db.add_all([existing, empty])
        await db.commit()

        cf_a = CandidateFile(
            tenant_id=tenant.id,
            candidate_id=existing.id,
            original_filename="a.pdf",
            mime_type="application/pdf",
            file_size=100,
            parse_status="pending",
            raw_text="resume text",
        )
        cf_b = CandidateFile(
            tenant_id=tenant.id,
            candidate_id=empty.id,
            original_filename="b.pdf",
            mime_type="application/pdf",
            file_size=100,
            parse_status="pending",
            raw_text="resume text",
        )
        db.add_all([cf_a, cf_b])
        await db.commit()

        monkeypatch.setattr(
            "app.modules.recruitment.tasks.parsing._extract_text_from_s3",
            lambda *args, **kwargs: "text",
        )
        captured = {"first_name": "X", "last_name": "Y"}

        async def fake_parse(_, **kwargs):
            return dict(captured)

        monkeypatch.setattr(
            "app.modules.recruitment.ai_service.parse_resume_text", fake_parse
        )

        retries: list[Exception] = []
        self._install_retry_stub(monkeypatch, retries)
        await _run_parse_task(str(cf_a.id), str(tenant.id))
        await _run_parse_task(str(cf_b.id), str(tenant.id))

        # Drop the in-memory views so the next ORM lookup re-reads the
        # rows the sync worker just committed.
        existing_id = existing.id
        empty_id = empty.id
        db.expunge_all()
        existing_row = (
            await db.execute(select(Candidate).where(Candidate.id == existing_id))
        ).scalar_one()
        empty_row = (
            await db.execute(select(Candidate).where(Candidate.id == empty_id))
        ).scalar_one()
        assert existing_row.parsed_resume_jsonb == {"existing": True}
        assert empty_row.parsed_resume_jsonb == captured

    async def test_failure_marks_row_and_retries(
        self, db: AsyncSession, tenant, user, monkeypatch
    ):
        cf = CandidateFile(
            tenant_id=tenant.id,
            candidate_id=None,
            original_filename="x.pdf",
            mime_type="application/pdf",
            file_size=10,
            parse_status="pending",
            raw_text="resume text",
        )
        db.add(cf)
        await db.commit()

        monkeypatch.setattr(
            "app.modules.recruitment.tasks.parsing._extract_text_from_s3",
            lambda *args, **kwargs: "text",
        )

        async def boom(_, **kwargs):
            raise RuntimeError("LLM blew up")

        monkeypatch.setattr(
            "app.modules.recruitment.ai_service.parse_resume_text", boom
        )

        retries: list[Exception] = []
        self._install_retry_stub(monkeypatch, retries)
        with pytest.raises(RuntimeError):
            await _run_parse_task(str(cf.id), str(tenant.id))
        assert len(retries) == 1

        await db.refresh(cf)
        assert cf.parse_status == "failed"
        assert "LLM blew up" in cf.parsed_data.get("error", "")


# ---------------------------------------------------------------------------
# Parsing-status + dedup-preview
# ---------------------------------------------------------------------------


class TestParsingStatus:
    async def test_returns_counts_and_per_file_preview(
        self, db: AsyncSession, tenant, user
    ):
        from app.modules.storage.models import File

        vacancy = await _make_vacancy(db, tenant.id, user.id)
        # Sweep C4: the empty-file_ids poll now scopes by File.uploaded_by,
        # so seed real File rows owned by the current user and wire them
        # to the CandidateFile rows through file_id.
        file_a = File(
            tenant_id=tenant.id,
            uploaded_by=user.id,
            name="a.pdf",
            original_name="a.pdf",
            mime_type="application/pdf",
            size=10,
            path="resumes/a.pdf",
        )
        file_b = File(
            tenant_id=tenant.id,
            uploaded_by=user.id,
            name="b.pdf",
            original_name="b.pdf",
            mime_type="application/pdf",
            size=10,
            path="resumes/b.pdf",
        )
        db.add_all([file_a, file_b])
        await db.flush()

        completed = CandidateFile(
            tenant_id=tenant.id,
            candidate_id=None,
            file_id=file_a.id,
            original_filename="a.pdf",
            mime_type="application/pdf",
            file_size=10,
            parse_status="completed",
            parsed_data={
                "first_name": "Ada",
                "last_name": "L",
                "contacts": {"email": "ada@example.com"},
                "experience": [{"position": "Engineer"}],
            },
        )
        pending = CandidateFile(
            tenant_id=tenant.id,
            candidate_id=None,
            file_id=file_b.id,
            original_filename="b.pdf",
            mime_type="application/pdf",
            file_size=10,
            parse_status="pending",
        )
        db.add_all([completed, pending])
        await db.commit()

        payload = await service.get_resumes_parsing_status(
            db, tenant.id, user.id, vacancy["id"], None
        )
        assert payload["counts"]["completed"] == 1
        assert payload["counts"]["pending"] == 1
        by_id = {item["file_id"]: item for item in payload["files"]}
        assert by_id[completed.id]["full_name"] == "Ada L"
        assert by_id[completed.id]["email"] == "ada@example.com"
        assert by_id[completed.id]["last_position"] == "Engineer"
        assert by_id[pending.id]["full_name"] is None


class TestDedupPreview:
    async def test_matches_active_ignores_archived(
        self, db: AsyncSession, tenant, user
    ):
        vacancy = await _make_vacancy(db, tenant.id, user.id)

        active = Candidate(
            tenant_id=tenant.id,
            full_name="Active Twin",
            email="dup@example.com",
        )
        archived = Candidate(
            tenant_id=tenant.id,
            full_name="Archived Twin",
            email="ghost@example.com",
            archived_at=datetime.now(timezone.utc) - timedelta(days=1),
        )
        db.add_all([active, archived])
        await db.flush()

        match_file = CandidateFile(
            tenant_id=tenant.id,
            candidate_id=None,
            original_filename="m.pdf",
            mime_type="application/pdf",
            file_size=10,
            parse_status="completed",
            parsed_data={"contacts": {"email": "DUP@example.com"}},
        )
        ghost_file = CandidateFile(
            tenant_id=tenant.id,
            candidate_id=None,
            original_filename="g.pdf",
            mime_type="application/pdf",
            file_size=10,
            parse_status="completed",
            parsed_data={"contacts": {"email": "ghost@example.com"}},
        )
        db.add_all([match_file, ghost_file])
        await db.commit()

        preview = await service.get_resumes_dedup_preview(
            db, tenant.id, vacancy["id"], [match_file.id, ghost_file.id]
        )
        by_id = {item["file_id"]: item for item in preview}
        assert by_id[match_file.id]["existing_candidate_id"] == active.id
        assert by_id[ghost_file.id]["existing_candidate_id"] is None


# ---------------------------------------------------------------------------
# Finalize uses Stage 3 LLM additions to populate denormalised fields
# ---------------------------------------------------------------------------


class TestFinalizeDenormalisation:
    async def test_creates_candidate_with_denormalised_fields(
        self, db: AsyncSession, tenant, user
    ):
        vacancy = await _make_vacancy(db, tenant.id, user.id)
        cf = CandidateFile(
            tenant_id=tenant.id,
            candidate_id=None,
            original_filename="p.pdf",
            mime_type="application/pdf",
            file_size=10,
            parse_status="completed",
            parsed_data={
                "first_name": "Priya",
                "last_name": "S",
                "current_position": "Engineering Manager",
                "years_of_experience": 11,
                "location": "Lisbon",
                "contacts": {
                    "email": "priya@example.com",
                    "phone": "+351 000",
                    "linkedin": "https://linkedin.com/in/priya",
                    "location": "Lisbon, PT",
                },
                "experience": [
                    {"position": "Engineering Manager", "role": "Engineering Manager"}
                ],
            },
        )
        db.add(cf)
        await db.commit()

        await service.finalize_candidates_from_parsed(
            db,
            tenant.id,
            user.id,
            vacancy["id"],
            BatchFinalizeRequest(files=[BatchFinalizeFile(file_id=cf.id)]),
        )

        cand = (
            await db.execute(
                select(Candidate).where(
                    Candidate.tenant_id == tenant.id,
                    Candidate.email == "priya@example.com",
                )
            )
        ).scalar_one()
        assert cand.current_position == "Engineering Manager"
        assert cand.years_of_experience == 11
        assert cand.location == "Lisbon"
        assert cand.linkedin_url == "https://linkedin.com/in/priya"

        cv = (
            await db.execute(
                select(CandidateVacancy).where(CandidateVacancy.candidate_id == cand.id)
            )
        ).scalar_one()
        assert cv.vacancy_id == vacancy["id"]


# ---------------------------------------------------------------------------
# LLM schema mapping
# ---------------------------------------------------------------------------


class TestLLMSchema:
    async def test_normalise_backfills_missing_top_level(self):
        payload = {
            "first_name": "Ravi",
            "last_name": "M",
            "contacts": {"location": "Bengaluru"},
            "experience": [
                {"role": "Director of Engineering"},
                {"role": "VP"},
            ],
        }
        ai_service._normalise_resume_payload(payload)
        assert payload["current_position"] == "Director of Engineering"
        assert payload["location"] == "Bengaluru"
        # ``position`` alias copied across so finalize/_last_position_from_parsed
        # see both keys uniformly.
        assert payload["experience"][0]["position"] == "Director of Engineering"
        assert payload["experience"][0]["role"] == "Director of Engineering"

    async def test_prompt_template_contains_new_fields(self):
        # Smoke check — the prompt must mention the Stage 3 additions so a
        # provider that respects the schema returns them.
        assert "current_position" in PARSE_RESUME
        assert "years_of_experience" in PARSE_RESUME
        assert "linkedin" in PARSE_RESUME

    async def test_prompt_experience_matches_card_contract(self):
        # HRP-346: the card renders start_date/end_date/description; the
        # prompt used to ask for period/achievements instead, so the
        # fields never reached the UI.
        assert "start_date" in PARSE_RESUME
        assert "end_date" in PARSE_RESUME
        assert '"description"' in PARSE_RESUME
        # Legacy JSON keys must be gone from the schema block.
        assert '"period"' not in PARSE_RESUME
        assert '"achievements"' not in PARSE_RESUME

    async def test_normalise_maps_legacy_experience_keys(self):
        # HRP-346: providers may still emit the pre-HRP-346 keys — split
        # ``period`` and copy ``achievements`` instead of dropping them.
        payload = {
            "experience": [
                {
                    "position": "QA Lead",
                    "company": "Acme",
                    "period": "2020 — 2023",
                    "achievements": "Built the QA team",
                },
                {
                    "role": "Engineer",
                    "period": "Mar 2018 to Feb 2020",
                    "achievements": ["Shipped v1", "Cut bugs by 40%"],
                },
                {
                    "position": "Intern",
                    "start_date": "2016",
                    "end_date": "2017",
                    "description": "kept as-is",
                    "period": "ignored",
                    "achievements": "ignored",
                },
            ]
        }
        ai_service._normalise_resume_payload(payload)
        first, second, third = payload["experience"]
        assert first["start_date"] == "2020"
        assert first["end_date"] == "2023"
        assert first["description"] == "Built the QA team"
        assert second["start_date"] == "Mar 2018"
        assert second["end_date"] == "Feb 2020"
        assert second["description"] == "Shipped v1\nCut bugs by 40%"
        # Explicit new-style keys win over legacy ones.
        assert third["start_date"] == "2016"
        assert third["end_date"] == "2017"
        assert third["description"] == "kept as-is"

    async def test_split_period_edge_cases(self):
        # Review finding: a naive split broke month names containing "to"
        # and intra-date hyphens.
        cases = {
            "October 2020 - March 2021": ("October 2020", "March 2021"),
            "2020-01 - 2023-05": ("2020-01", "2023-05"),
            "2020-2023": ("2020", "2023"),
            "2020 — 2023": ("2020", "2023"),
            "Mar 2020 to Present": ("Mar 2020", "Present"),
            "Toronto 2020 - 2021": ("Toronto 2020", "2021"),
        }
        for period, (start, end) in cases.items():
            parts = ai_service._split_period(period)
            assert parts == [start, end], period
        # A single date must not be split.
        assert ai_service._split_period("2020") == ["2020"]
        assert ai_service._split_period("March 2021") == ["March 2021"]

    async def test_canonical_read_normalises_legacy_payload(self):
        # Candidates parsed before HRP-346 keep period/achievements in
        # stored JSONB — the read path maps them without mutating storage.
        from app.modules.recruitment.candidate_service import (
            _normalised_parsed_resume,
        )

        stored = {
            "experience": [
                {"position": "QA", "period": "2019 - 2021", "achievements": "Did X"}
            ]
        }
        out = _normalised_parsed_resume(stored)
        assert out["experience"][0]["start_date"] == "2019"
        assert out["experience"][0]["end_date"] == "2021"
        assert out["experience"][0]["description"] == "Did X"
        assert "start_date" not in stored["experience"][0]


# ---------------------------------------------------------------------------
# DOCX text extraction (HRP-345)
# ---------------------------------------------------------------------------


class TestDocxExtraction:
    def _build_docx(self) -> bytes:
        from docx import Document

        doc = Document()
        # Header content — contact lines often live here in templates.
        doc.sections[0].header.paragraphs[0].text = "Jane Doe — jane@example.com"
        # Two-column resume templates keep the whole body in a table;
        # ``doc.paragraphs`` sees nothing.
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Experience"
        table.cell(0, 1).text = "QA Engineer at Acme, 2020-2023"
        nested = table.cell(1, 1).add_table(rows=1, cols=1)
        nested.cell(0, 0).text = "Playwright, pytest"
        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    async def test_table_only_docx_extracts_text(self):
        # HRP-345: table-based DOCX resumes extracted to "" and the whole
        # upload surfaced as "failed parsing".
        from app.modules.recruitment.tasks.parsing import _extract_docx_text

        text = _extract_docx_text(self._build_docx())
        assert "QA Engineer at Acme, 2020-2023" in text
        assert "Playwright, pytest" in text
        assert "jane@example.com" in text

    async def test_paragraph_docx_still_extracts(self):
        from app.modules.recruitment.tasks.parsing import _extract_docx_text
        from docx import Document

        doc = Document()
        doc.add_paragraph("Plain paragraph resume")
        buf = BytesIO()
        doc.save(buf)
        assert "Plain paragraph resume" in _extract_docx_text(buf.getvalue())


pytestmark = pytest.mark.asyncio
